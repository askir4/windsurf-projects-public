"""
Report Generator Module
Handles the logic for generating reports by filling templates with data
"""

import re
import logging
from typing import Dict, Any, Optional
from datetime import datetime, date
from copy import deepcopy
from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates reports by replacing placeholders in templates with actual data."""
    
    PLACEHOLDER_PATTERN = re.compile(r'\{\{(\w+)\}\}')
    
    def __init__(self):
        """Initialize the report generator."""
        pass
    
    def _format_value(self, value: Any) -> str:
        """
        Format a value for insertion into the document.
        
        Args:
            value: The value to format
            
        Returns:
            Formatted string representation
        """
        if value is None:
            return ""
        elif isinstance(value, datetime):
            return value.strftime("%d.%m.%Y")
        elif isinstance(value, date):
            return value.strftime("%d.%m.%Y")
        elif isinstance(value, (int, float)):
            return str(value)
        elif isinstance(value, list):
            return "\n".join(str(item) for item in value)
        else:
            return str(value)
    
    def _replace_in_paragraph(self, paragraph, data: Dict[str, Any]) -> None:
        """
        Replace placeholders in a paragraph while preserving formatting.
        
        Args:
            paragraph: The paragraph to process
            data: Dictionary mapping placeholder names to values
        """
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Find all placeholders
        matches = list(self.PLACEHOLDER_PATTERN.finditer(full_text))
        
        if not matches:
            return
        
        # Build the new text with replacements
        new_text = full_text
        for match in reversed(matches):  # Reverse to maintain positions
            placeholder_name = match.group(1)
            if placeholder_name in data:
                replacement = self._format_value(data[placeholder_name])
                new_text = new_text[:match.start()] + replacement + new_text[match.end():]
        
        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear existing runs and add new text
            # We need to preserve the formatting of the first run
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Store formatting
                font_name = first_run.font.name
                font_size = first_run.font.size
                bold = first_run.font.bold
                italic = first_run.font.italic
                underline = first_run.font.underline
                
                # Clear all runs
                for run in paragraph.runs:
                    run.text = ""
                
                # Set new text on first run
                first_run.text = new_text
                
                # Restore formatting
                if font_name:
                    first_run.font.name = font_name
                if font_size:
                    first_run.font.size = font_size
                if bold is not None:
                    first_run.font.bold = bold
                if italic is not None:
                    first_run.font.italic = italic
                if underline is not None:
                    first_run.font.underline = underline
            else:
                paragraph.add_run(new_text)
    
    def _replace_in_table(self, table, data: Dict[str, Any]) -> None:
        """
        Replace placeholders in all cells of a table.
        
        Args:
            table: The table to process
            data: Dictionary mapping placeholder names to values
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_in_paragraph(paragraph, data)
    
    def _replace_in_header_footer(self, header_footer, data: Dict[str, Any]) -> None:
        """
        Replace placeholders in header or footer.
        
        Args:
            header_footer: The header or footer to process
            data: Dictionary mapping placeholder names to values
        """
        if header_footer:
            for paragraph in header_footer.paragraphs:
                self._replace_in_paragraph(paragraph, data)
            for table in header_footer.tables:
                self._replace_in_table(table, data)
    
    def generate(self, template_doc: Document, data: Dict[str, Any]) -> Document:
        """
        Generate a report by filling a template with data.
        
        Args:
            template_doc: The template Document object
            data: Dictionary mapping placeholder names to values
            
        Returns:
            A new Document with placeholders replaced
        """
        # Create a copy of the document to avoid modifying the original
        # Note: python-docx doesn't support deep copy, so we work with the original
        # The caller should pass a freshly loaded document
        doc = template_doc
        
        # Replace in all paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data)
        
        # Replace in all tables
        for table in doc.tables:
            self._replace_in_table(table, data)
        
        # Replace in headers and footers
        for section in doc.sections:
            self._replace_in_header_footer(section.header, data)
            self._replace_in_header_footer(section.footer, data)
        
        logger.info(f"Report generated with {len(data)} data fields")
        return doc
    
    def validate_data(self, placeholders: set, data: Dict[str, Any]) -> Dict[str, list]:
        """
        Validate that all required placeholders have data.
        
        Args:
            placeholders: Set of placeholder names from the template
            data: Dictionary of data to be inserted
            
        Returns:
            Dictionary with 'missing' and 'extra' lists
        """
        data_keys = set(data.keys())
        
        missing = placeholders - data_keys
        extra = data_keys - placeholders
        
        return {
            'missing': sorted(missing),
            'extra': sorted(extra)
        }
    
    def get_default_data(self) -> Dict[str, Any]:
        """
        Get default data fields commonly used in training reports.
        
        Returns:
            Dictionary with default field names and empty/default values
        """
        today = date.today()
        
        return {
            'name': '',
            'vorname': '',
            'nachname': '',
            'datum': today,
            'datum_von': today,
            'datum_bis': today,
            'woche': today.isocalendar()[1],
            'jahr': today.year,
            'abteilung': '',
            'ausbilder': '',
            'betrieb': '',
            'beruf': '',
            'ausbildungsjahr': 1,
            'taetigkeiten': '',
            'betriebliche_taetigkeiten': '',
            'unterweisungen': '',
            'berufsschule': '',
            'stunden': 40,
            'stunden_betrieb': 0,
            'stunden_schule': 0,
            'unterschrift_azubi': '',
            'unterschrift_ausbilder': '',
            'bemerkungen': ''
        }
