"""
Template Handler Module
Manages loading, saving, and analyzing Word templates
"""

import os
import re
import shutil
import logging
from pathlib import Path
from typing import List, Set, Optional
from docx import Document

logger = logging.getLogger(__name__)


class TemplateHandler:
    """Handles Word template operations including loading, saving, and placeholder detection."""
    
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([^}]+)\}\}')
    
    def __init__(self, templates_dir: str = "templates"):
        """
        Initialize the template handler.
        
        Args:
            templates_dir: Directory path for storing templates
        """
        self.templates_dir = Path(templates_dir)
        self._ensure_templates_dir()
    
    def _ensure_templates_dir(self) -> None:
        """Create templates directory if it doesn't exist."""
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"Templates directory ensured at: {self.templates_dir}")
    
    def import_template(self, source_path: str, template_name: Optional[str] = None) -> str:
        """
        Import a Word template file into the templates directory.
        
        Args:
            source_path: Path to the source .docx file
            template_name: Optional custom name for the template
            
        Returns:
            The filename of the imported template
            
        Raises:
            FileNotFoundError: If source file doesn't exist
            ValueError: If file is not a .docx file
        """
        source = Path(source_path)
        
        if not source.exists():
            raise FileNotFoundError(f"Template file not found: {source_path}")
        
        if source.suffix.lower() != '.docx':
            raise ValueError(f"Invalid file format. Expected .docx, got: {source.suffix}")
        
        # Determine destination filename
        if template_name:
            if not template_name.endswith('.docx'):
                template_name += '.docx'
            dest_name = template_name
        else:
            dest_name = source.name
        
        dest_path = self.templates_dir / dest_name
        
        # Copy file to templates directory
        shutil.copy2(source, dest_path)
        logger.info(f"Template imported: {dest_name}")
        
        return dest_name
    
    def list_templates(self) -> List[str]:
        """
        List all available templates in the templates directory.
        
        Returns:
            List of template filenames
        """
        templates = [f.name for f in self.templates_dir.glob("*.docx") if f.is_file()]
        templates.sort()
        return templates
    
    def get_template_path(self, template_name: str) -> Path:
        """
        Get the full path to a template file.
        
        Args:
            template_name: Name of the template file
            
        Returns:
            Full path to the template
            
        Raises:
            FileNotFoundError: If template doesn't exist
        """
        template_path = self.templates_dir / template_name
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_name}")
        return template_path
    
    def delete_template(self, template_name: str) -> bool:
        """
        Delete a template from the templates directory.
        
        Args:
            template_name: Name of the template to delete
            
        Returns:
            True if deletion was successful
            
        Raises:
            FileNotFoundError: If template doesn't exist
        """
        template_path = self.get_template_path(template_name)
        template_path.unlink()
        logger.info(f"Template deleted: {template_name}")
        return True
    
    def load_template(self, template_name: str) -> Document:
        """
        Load a Word template as a Document object.
        
        Args:
            template_name: Name of the template file
            
        Returns:
            python-docx Document object
        """
        template_path = self.get_template_path(template_name)
        return Document(template_path)
    
    def extract_placeholders(self, template_name: str) -> Set[str]:
        """
        Extract all placeholders from a template.
        
        Placeholders are in the format {{placeholder_name}}.
        
        Args:
            template_name: Name of the template file
            
        Returns:
            Set of placeholder names (without braces)
        """
        doc = self.load_template(template_name)
        placeholders = set()
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
            placeholders.update(matches)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                        placeholders.update(matches)
        
        # Search in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                    placeholders.update(matches)
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    matches = self.PLACEHOLDER_PATTERN.findall(paragraph.text)
                    placeholders.update(matches)
        
        logger.info(f"Found {len(placeholders)} placeholders in {template_name}: {placeholders}")
        return placeholders
    
    def get_template_info(self, template_name: str) -> dict:
        """
        Get information about a template.
        
        Args:
            template_name: Name of the template file
            
        Returns:
            Dictionary with template information
        """
        template_path = self.get_template_path(template_name)
        placeholders = self.extract_placeholders(template_name)
        
        stat = template_path.stat()
        
        return {
            'name': template_name,
            'path': str(template_path),
            'size_bytes': stat.st_size,
            'modified': stat.st_mtime,
            'placeholders': sorted(placeholders),
            'placeholder_count': len(placeholders)
        }
